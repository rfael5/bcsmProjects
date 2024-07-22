import re
from tkinter import Button, Text, ttk
import connection
import tabela

import tkinter as tk
from tkinter import Entry, Label, StringVar
from docx import Document


class FilterableCombobox(tk.Frame):
    def __init__(self, parent, options, tabela, dados_evento = None, entry_width=40, materiais_evento = []):
        super().__init__(parent)
        self.options = options
        self.tabela = tabela
        self.var = StringVar()
        self.materiais_evento = materiais_evento
        
        self.dados_evento = dados_evento
                
        # Entry widget
        self.entry = tk.Entry(self, textvariable=self.var, width=entry_width)
        self.entry.grid(row=1, column=0, padx=(0, 100), columnspan=2, sticky='nsew')
        
        # Listbox widget
        self.listbox = tk.Listbox(self, width=entry_width)
        self.listbox.grid(row=2, column=0, sticky="ew")
        self.listbox.grid_remove()  # Esconde o Listbox inicialmente
        
        self.qtd_material = StringVar()
        label_qtd = Label(self, text="Quantidade:", font=('Arial', 12, 'bold'))
        label_qtd.grid(row=3, column=0)

        input_qtd = Entry(self, textvariable=self.qtd_material, width=20)
        input_qtd.grid(row=4, column=0, padx=(0,200), columnspan=2, sticky='nsew')
        
        self.update_listbox()

        # Bindings
        self.entry.bind('<KeyRelease>', self.on_keyrelease)
        self.listbox.bind('<<ListboxSelect>>', self.on_select)

        # Configurar grid para expandir corretamente
        self.grid_columnconfigure(0, weight=1)
    
    def appendMaterial(self):
        _material = self.var.get()
        _quantidade = self.qtd_material.get()
        materiais = {
            "material": _material,
            "quantidade": _quantidade 
        }
        print(materiais)
        self.materiais_evento.append(materiais)
        
        #for m in lista_materiais:
        produto = materiais['material']
        quantidade = materiais['quantidade']
        data = (produto, quantidade)
        self.tabela.insert(parent='', index=0, values=data)
        
        self.var.set("")
        self.qtd_material.set("")
    
    def ajustarString(self, texto):
        if '\x00' in texto:
            clear_line = texto.replace('\x00', '')
            encoded_line = clear_line.encode('utf-8')
            return encoded_line
        else:
            encoded_line = texto.encode('utf-8')
            return encoded_line
    
    def formatarStringDict(self, dicionario):
        dicionario['NOME'] = self.ajustarString(dicionario['NOME'])
        dicionario['DOCUMENTO'] = self.ajustarString(str(dicionario['DOCUMENTO']))
        dicionario['DESCRICAO'] = self.ajustarString(dicionario['DESCRICAO'])
        dicionario['DTEVENTO'] = self.ajustarString(str(dicionario['DTEVENTO']))
        dicionario['CONVIDADOS'] = self.ajustarString(str(dicionario['CONVIDADOS']))
        dicionario['NOMEINTERNO'] = self.ajustarString(dicionario['NOMEINTERNO'])
        dicionario['LOCAL'] = self.ajustarString(dicionario['LOCAL'])
        
    
    def criarDocumentoWord(self, lista):
        doc = Document()
        
        self.formatarStringDict(self.dados_evento)
        
        doc.add_heading(f'{self.dados_evento["NOME"].decode("utf-8")}         OR {self.dados_evento["DOCUMENTO"].decode("utf-8")}')
        tabela_info = doc.add_table(rows=3, cols=2)
        
        first_row = tabela_info.rows[0].cells
        first_row[0].text = f"Evento: Evento"
        first_row[1].text = f"Serviço: {self.dados_evento['DESCRICAO'].decode('utf-8')}"
        
        second_row = tabela_info.rows[1].cells
        #row_cells[0].text = f"Serviço: {self.dados_evento['DESCRICAO'].decode('utf-8')}"
        second_row[0].text = f"Data do evento: {self.dados_evento['DTEVENTO'].decode('utf-8')}"
        second_row[1].text = f"Convidados: {self.dados_evento['CONVIDADOS'].decode('utf-8')}"
        
        third_row = tabela_info.rows[2].cells
        third_row[0].text = f"Consultor: {self.dados_evento['NOMEINTERNO'].decode('utf-8')}"
        third_row[1].text = f"Local do evento: {self.dados_evento['LOCAL'].decode('utf-8')}"
        
        
        for material in lista:
            str_material = self.ajustarString(material['material'])
            #print(encoded_line)
            doc.add_paragraph(f'''{str_material.decode('utf-8')} - {material['quantidade']} unidades''')
        doc.save('lista-materiais.docx')

    def on_keyrelease(self, event):
        value = event.widget.get()
        value = value.strip().lower()
        
        if value == '':
            self.listbox.grid_remove()
        else:
            data = []
            for item in self.options:
                if value in item.lower():
                    data.append(item)
            self.update_listbox(data)
    
    def on_select(self, event):
        # Verificar se há uma seleção válida na Listbox
        try:
            selection = self.listbox.curselection()
            if selection:
                selected_item = self.listbox.get(selection)
                self.var.set(selected_item)
                self.listbox.grid_remove()
        except tk.TclError:
            pass
        
    def update_listbox(self, data=None):
        if data is None:
            data = self.options
        
        self.listbox.delete(0, 'end')

        for item in data:
            self.listbox.insert('end', item)
        
        if data:
            self.listbox.grid()
        else:
            self.listbox.grid_remove()

def criarListaMateriais():
    materiais = connection.getMateriais()
    lista_materiais = []
    for produto in materiais:
        lista_materiais.append(produto['DESCRICAO'])
    return lista_materiais



def setarDatosEvento():
    ...

def main():
    root = tk.Tk()
    root.title("Digitalizar lista de materiais")

    root.geometry("1150x800")

    notebook = ttk.Notebook(root)
    notebook.pack(fill=tk.BOTH, expand=True)

    page1 = tk.Frame(notebook)
    notebook.add(page1, text='| |')

    mainFrame = tk.Frame(page1)
    mainFrame.pack(fill=tk.BOTH, expand=1)

    canvas = tk.Canvas(mainFrame)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=1)
    #canvas.grid(row=0, column=0, sticky=EW)

    scrollbar = ttk.Scrollbar(mainFrame, orient=tk.VERTICAL, command=canvas.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    #scrollbar.grid(row=0, rowspan=10, column=1, sticky="ns")

    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.bind('<Configure>', lambda e:canvas.configure(scrollregion=canvas.bbox("all")))

    secondFrame = tk.Frame(canvas)
    canvas.create_window((0, 0), window=secondFrame, anchor="nw")

    options = criarListaMateriais()
    
    tabela.criarTabela(secondFrame)
    
    info_evento = None
    
    num_or = StringVar()
    label_numevento = Label(secondFrame, text="Número OR:", font=('Arial', 12, 'bold'))
    label_numevento.grid(row=0, column=0)
    input_or = Entry(secondFrame, textvariable=num_or)
    input_or.grid(row=0, column=1, padx=2, sticky='s')
    btn_buscar_evento = Button(secondFrame, text="Buscar evento", bg='#C0C0C0', font=("Arial", 12), command= lambda: buscarEvento(num_or, info_evento))
    btn_buscar_evento.grid(row=0, column=2)
    
    lbl_or = Label(secondFrame, text="OR: ", font=('Arial', 12))
    lbl_or.grid(row=1, column=0)
    
    lbl_dtevento = Label(secondFrame, text="Data evento: ", font=('Arial', 12))
    lbl_dtevento.grid(row=2, column=0)
    
    lbl_evento = Label(secondFrame, text="Evento: ", font=('Arial', 12))
    lbl_evento.grid(row=3, column=0)
    
    lbl_convidados = Label(secondFrame, text="Convidados: ", font=('Arial', 12))
    lbl_convidados.grid(row=4, column=0)
    
    lbl_servico = Label(secondFrame, text="Serviço: ", font=('Arial', 12))
    lbl_servico.grid(row=5, column=0)
    
    lbl_consultor = Label(secondFrame, text="Consultor(a): ", font=('Arial', 12))
    lbl_consultor.grid(row=6, column=0)
    
    lbl_local = Label(secondFrame, text="Local do evento: ", font=('Arial', 12))
    lbl_local.grid(row=7, column=0)
    
    lbl_nome = Label(secondFrame, text="Nome: ", font=('Arial', 12))
    lbl_nome.grid(row=8, column=0)
    
    material = ''
    label_qtd = Label(secondFrame, text="Produto:", font=('Arial', 12, 'bold'))
    label_qtd.grid(row=9, column=0)
    combobox = FilterableCombobox(secondFrame, options, tabela.table, info_evento)
    combobox.grid(row=10, column=0, padx=(0, 100), columnspan=2, sticky='nsew')
    
    btn = Button(secondFrame, text="Adicionar a lista", bg='#C0C0C0', font=("Arial", 16), command=combobox.appendMaterial)
    btn.grid(row=11, column=0, pady=20)
    
    btn_doc = Button(secondFrame, text="Gerar documento", bg='#C0C0C0', font=("Arial", 16), command=lambda: combobox.criarDocumentoWord(combobox.materiais_evento))
    btn_doc.grid(row=13, column=0, pady=20)
    
    def buscarEvento(n, info):
        documento = n.get()
        evento = connection.getEvento(documento)
        combobox.dados_evento = evento[0]
        
        lbl_nome.config(text=f"Nome: {evento[0]['NOME']}")
        lbl_local.config(text=f"Local do evento: {evento[0]['LOCAL']}")
        lbl_consultor.config(text=f"Consultor(a): {evento[0]['NOMEINTERNO']}")
        lbl_servico.config(text=f"Serviço: {evento[0]['DESCRICAO']}")
        lbl_convidados.config(text=f"Convidados: {evento[0]['CONVIDADOS']}")
        lbl_evento.config(text=f"Evento: Evento")
        lbl_dtevento.config(text=f"Data evento: {evento[0]['DTEVENTO']}")
        lbl_or.config(text=f"OR: {evento[0]['DOCUMENTO']}")

    root.mainloop()

if __name__ == "__main__":
    main()